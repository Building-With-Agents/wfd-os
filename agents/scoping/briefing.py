"""Generate the pre-call briefing Word document."""

import os
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.scoping.models import ScopingRequest, ResearchResult


def generate_briefing_doc(req: ScopingRequest, research: ResearchResult) -> str:
    """Create a briefing .docx and return the local file path."""
    org = req.organization
    contact = req.contact
    today = date.today()
    company = org.safe_name

    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -- Header --
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("COMPUTING FOR ALL")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(153, 0, 0)

    doc.add_paragraph()  # spacer

    # -- Title --
    title = doc.add_heading(f"Pre-Call Briefing: {org.name}", level=1)

    doc.add_paragraph(f"Prepared: {today.strftime('%B %d, %Y')}")
    doc.add_paragraph()

    # -- 1. Prospect Overview --
    doc.add_heading("1. Prospect Overview", level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Light Grid Accent 1"
    fields = [
        ("Company", org.name),
        ("Industry", org.industry or "—"),
        ("Size", org.employee_count or "—"),
        ("Website", org.website_url or "—"),
        ("Primary Contact", f"{contact.full_name}, {contact.title}"),
        ("Email", contact.email or "—"),
    ]
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # -- 2. Mission & Strategic Context --
    doc.add_heading("2. Mission & Strategic Context", level=2)
    doc.add_paragraph(research.mission_and_strategy or research.company_overview or "Research pending.")

    # -- 3. Recent News & Signals --
    doc.add_heading("3. Recent News & Signals", level=2)
    if research.recent_news:
        for item in research.recent_news:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No recent news found during research.")

    # -- 4. Technology & Data Landscape --
    doc.add_heading("4. Technology & Data Landscape", level=2)
    doc.add_paragraph(research.tech_landscape or "Limited information available — explore during scoping call.")

    # -- 5. Likely Pain Points --
    doc.add_heading("5. Likely Pain Points", level=2)
    if research.likely_pain_points:
        for point in research.likely_pain_points:
            doc.add_paragraph(point, style="List Bullet")
    else:
        doc.add_paragraph("To be explored during scoping call.")

    # -- 6. Suggested Scoping Questions --
    doc.add_heading("6. Suggested Scoping Questions", level=2)
    if research.suggested_questions:
        for i, q in enumerate(research.suggested_questions, 1):
            doc.add_paragraph(f"{i}. {q}")
    else:
        # Default questions
        defaults = [
            "What is the primary problem you are looking to solve with AI or data automation?",
            "What does your current data infrastructure look like?",
            "Who would be the internal champion and decision maker for this project?",
            "What does success look like for this engagement?",
            "What is your timeline and budget range for this work?",
        ]
        for i, q in enumerate(defaults, 1):
            doc.add_paragraph(f"{i}. {q}")

    # -- 7. Contact Details --
    doc.add_heading("7. Contact Details", level=2)
    doc.add_paragraph(f"Name: {contact.full_name}")
    doc.add_paragraph(f"Title: {contact.title}")
    doc.add_paragraph(f"Email: {contact.email}")
    if contact.linkedin_url:
        doc.add_paragraph(f"LinkedIn: {contact.linkedin_url}")

    # -- Sales notes --
    if req.notes:
        doc.add_heading("Sales Notes (from Jason)", level=2)
        doc.add_paragraph(req.notes)

    # -- Footer on all pages --
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential — Computing for All | computingforall.org")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    os.makedirs("output", exist_ok=True)
    filename = f"Briefing_{company}_{today.isoformat()}.docx"
    filepath = os.path.join("output", filename)
    doc.save(filepath)
    print(f"[DOCS] Briefing doc saved to {filepath}")
    return filepath
