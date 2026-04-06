"""Generate the proposal Word document after scoping call analysis."""

import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.scoping.models import ScopingRequest, ScopingAnalysis


def generate_proposal_doc(req: ScopingRequest, analysis: ScopingAnalysis) -> str:
    """Create a proposal .docx and return the local file path."""
    org = req.organization
    contact = req.contact
    today = date.today()
    company = org.safe_name

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- 1. Cover Page ----
    for _ in range(4):
        doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("COMPUTING FOR ALL")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    proposal_title = doc.add_paragraph()
    proposal_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proposal_title.add_run(f"Proposal: AI Agent Engagement")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    for text in [
        f"Prepared for: {org.name}",
        "Prepared by: Computing for All",
        f"Date: {today.strftime('%B %d, %Y')}",
        "",
        "CONFIDENTIAL",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        if text == "CONFIDENTIAL":
            run.font.color.rgb = RGBColor(153, 0, 0)
            run.bold = True

    doc.add_page_break()

    # ---- 2. Executive Summary ----
    doc.add_heading("Executive Summary", level=1)

    problem_text = analysis.problem_summary or "[Problem summary — to be completed from scoping analysis]"
    doc.add_paragraph(
        f"The Problem: {problem_text}"
    )
    doc.add_paragraph(
        "Proposed Solution: CFA will design and build a custom AI agent system "
        f"tailored to {org.name}'s specific operational needs. The solution will "
        "automate key workflows, reduce manual effort, and deliver measurable "
        "efficiency gains."
    )
    doc.add_paragraph(
        "Why CFA: Computing for All is an agentic data engineering firm that combines "
        "cutting-edge AI agent architecture with a supervised apprenticeship delivery model. "
        "Our teams build production-grade agent systems while developing the next generation "
        "of AI engineers."
    )
    doc.add_paragraph(
        "Investment: [RITU TO COMPLETE]"
    )

    # ---- 3. The Problem ----
    doc.add_heading("The Problem", level=1)

    # Pull from transcript analysis if available
    problem_answer = ""
    if analysis.answers:
        problem_answer = analysis.answers[0].answer  # Q1 = problem
    doc.add_paragraph(
        problem_answer or "[Detailed problem statement drawn from scoping call transcript — "
        "to be completed from analysis]"
    )
    doc.add_paragraph(
        f"Current state: {org.name} faces operational friction in areas where "
        "manual processes, disconnected data, and lack of automation create "
        "bottlenecks. [GARY TO COMPLETE with technical specifics]"
    )
    doc.add_paragraph(
        "When solved: The organization will have an automated, reliable system "
        "that reduces manual work, improves data accuracy, and enables faster "
        "decision-making."
    )

    # ---- 4. Proposed Solution ----
    doc.add_heading("Proposed Solution", level=1)
    doc.add_paragraph("[GARY TO COMPLETE]")
    doc.add_paragraph(
        "CFA will build a custom AI agent system that addresses the core problem "
        f"identified during our scoping conversation with {contact.full_name}. "
        "The solution will leverage CFA's Data Unlock framework to identify the "
        "highest-impact automation opportunities."
    )
    doc.add_paragraph("Technical approach: [GARY TO COMPLETE]")
    doc.add_paragraph("Deliverables at completion: [GARY TO COMPLETE]")

    # ---- 5. Scope and Deliverables ----
    doc.add_heading("Scope and Deliverables", level=1)
    doc.add_paragraph("Deliverable 1: [GARY TO COMPLETE]")
    doc.add_paragraph("Deliverable 2: [GARY TO COMPLETE]")
    doc.add_paragraph("Deliverable 3: [GARY TO COMPLETE]")
    doc.add_paragraph()
    doc.add_paragraph("Out of scope: [GARY TO COMPLETE]")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: Scope to be confirmed with Gary prior to proposal submission.")
    run.italic = True

    # ---- 6. Engagement Timeline ----
    doc.add_heading("Engagement Timeline", level=1)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Phase", "Activities", "Owner", "Timeline"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    phases = [
        ("Phase 1: Scoping & Setup", "[GARY TO COMPLETE]", "CFA + Client", "Week 1–2"),
        ("Phase 2: Build", "[GARY TO COMPLETE]", "CFA Engineering", "Week 3–8"),
        ("Phase 3: Deploy & Handover", "[GARY TO COMPLETE]", "CFA + Client", "Week 8"),
    ]
    for i, (phase, activities, owner, timeline) in enumerate(phases, 1):
        table.rows[i].cells[0].text = phase
        table.rows[i].cells[1].text = activities
        table.rows[i].cells[2].text = owner
        table.rows[i].cells[3].text = timeline

    # ---- 7. The CFA Team ----
    doc.add_heading("The CFA Team", level=1)
    doc.add_paragraph(
        "Ritu Bahl — Executive Director. Ritu oversees all CFA engagements and "
        "ensures alignment between client needs and delivery."
    )
    doc.add_paragraph(
        "Gary — Technical Lead. Gary leads technical architecture, solution design, "
        "and engineering delivery across all CFA projects."
    )
    doc.add_paragraph(
        "Apprentice Engineering Team — CFA's delivery model pairs experienced engineers "
        "with apprentice developers in supervised teams. All apprentice work is reviewed "
        "and validated by Gary before delivery. This model produces high-quality output "
        "while creating pathways into technology careers for underrepresented talent."
    )

    # ---- 8. Investment ----
    doc.add_heading("Investment", level=1)
    doc.add_paragraph("Consulting fee: [RITU TO COMPLETE]")
    doc.add_paragraph("Managed services (optional): [RITU TO COMPLETE]")
    doc.add_paragraph("Payment schedule: 50% at signing, 50% at delivery (standard)")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Placement fee note: If the client elects to hire an apprentice from the "
        "CFA team, a separate placement agreement will apply."
    )
    run.italic = True

    # ---- 9. Next Steps ----
    doc.add_heading("Next Steps", level=1)
    doc.add_paragraph(f"1. Review and confirm scope with Gary")
    doc.add_paragraph(f"2. Ritu to finalize investment figures")
    doc.add_paragraph(f"3. Jason to present to {contact.full_name}")
    doc.add_paragraph(f"4. Contract via DocuSeal upon agreement")

    # ---- Gaps section (internal — for team only) ----
    if analysis.gaps or analysis.followup_questions:
        doc.add_page_break()
        doc.add_heading("INTERNAL — Scoping Gaps (Do Not Send to Client)", level=1)
        p = doc.add_paragraph()
        run = p.add_run("The following items need follow-up before this proposal is finalized:")
        run.italic = True
        run.font.color.rgb = RGBColor(153, 0, 0)

        if analysis.gaps:
            doc.add_heading("Gaps Identified", level=2)
            for gap in analysis.gaps:
                doc.add_paragraph(gap, style="List Bullet")

        if analysis.followup_questions:
            doc.add_heading("Suggested Follow-Up Questions", level=2)
            for q in analysis.followup_questions:
                doc.add_paragraph(q, style="List Bullet")

    # ---- Footer on all pages ----
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential — Computing for All | computingforall.org")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    os.makedirs("output", exist_ok=True)
    filename = f"Proposal_{company}_{today.isoformat()}.docx"
    filepath = os.path.join("output", filename)
    doc.save(filepath)
    print(f"[DOCS] Proposal doc saved to {filepath}")
    return filepath
